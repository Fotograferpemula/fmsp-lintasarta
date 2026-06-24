import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_bottom_border(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(23, 105, 255) # Lintasarta Blue
        run.bold = True
        
        pPr = heading._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # 1.5 pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1769FF')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(13, 79, 204) # Darker Blue
        run.bold = True
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(60, 60, 60)
        run.bold = True
        run.italic = True
        
    return heading

def build_table(doc, headers, rows_data, widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "1769FF")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    # Rows
    for idx, row_data in enumerate(rows_data):
        row_cells = table.rows[idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=150, right=150)
            if idx % 2 == 1:
                set_cell_shading(row_cells[c_idx], "F9FAFB")
            p = row_cells[c_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if c_idx == 0:
                p.runs[0].bold = True
                
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
                
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    return table

def generate_detailed_document():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    logo_path = "public/docs/screenshots/logo.png"
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(30)
        logo_run = logo_p.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.8))

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_run = title_p.add_run("BUSINESS REQUIREMENTS DOCUMENT (BRD)\nSOFTWARE REQUIREMENTS SPECIFICATION (SRS)\n& SPESIFIKASI ARSITEKTUR FULL STACK")
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(23, 105, 255)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    sub_run = sub_p.add_run("FACILITY MANAGEMENT SERVICE PLATFORM (FMSP) LINTASARTA")
    sub_run.font.size = Pt(13)
    sub_run.bold = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run("―" * 45)
    line_run.font.color.rgb = RGBColor(23, 105, 255)
    
    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p.paragraph_format.space_before = Pt(40)
    version_run = version_p.add_run(
        "Dokumen Standar Spesifikasi Kebutuhan & Desain Teknis Sistem Terintegrasi\n"
        "Nomor Dokumen: LA-FMSP-BRD-SRS-001\n"
        "Kategori: Dokumen Teknis Rahasia\n"
        "Versi 2.2 - Juni 2026"
    )
    version_run.font.size = Pt(10.5)
    version_run.font.color.rgb = RGBColor(120, 120, 120)
    
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.paragraph_format.space_before = Pt(80)
    org_run = org_p.add_run(
        "PT Aplikanusa Lintasarta\n"
        "Divisi Facility Management & Keamanan Informasi\n"
        "Departemen IT Infrastructure & DevOps"
    )
    org_run.bold = True
    org_run.font.size = Pt(12)
    org_run.font.color.rgb = RGBColor(23, 105, 255)
    
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # DOCUMENT CONTROL
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Document Control & History", level=1)
    
    doc.add_paragraph(
        "Dokumen ini memuat detail rencana bisnis, spesifikasi fungsional, non-fungsional, serta "
        "arsitektur teknis dari platform FMSP Lintasarta. Setiap perubahan pada dokumen ini harus dicatat pada tabel di bawah ini:"
    )
    
    headers_ctrl = ["Versi", "Tanggal", "Penulis", "Deskripsi Perubahan"]
    rows_ctrl = [
        ("1.0", "15 April 2026", "GA Team", "Inisiasi awal dokumen BRD Fase 1."),
        ("2.0", "20 Mei 2026", "DevOps Team", "Penambahan kebutuhan teknis Fase 2 (RAB, Keuangan, WO)."),
        ("2.1", "22 Juni 2026", "Auditor / Zed", "Audit kesesuaian sistem dan perbaikan inkonsistensi RBAC."),
        ("2.2", "24 Juni 2026", "Antigravity", "Penyempurnaan detail standar BRD & SRS, spesifikasi schema lengkap, API, dan skema HA.")
    ]
    build_table(doc, headers_ctrl, rows_ctrl, [0.8, 1.2, 1.5, 3.0])
    
    doc.add_paragraph("Daftar Penyetuju Dokumen (Approvals):")
    headers_appr = ["Nama / Peran", "Jabatan", "Tanggal Setuju", "Tanda Tangan Status"]
    rows_appr = [
        ("Bambang W.", "General Manager GA Lintasarta", "24 Juni 2026", "APPROVED (Sistem Digital)"),
        ("Rian H.", "Head of IT Infrastructure Lintasarta", "24 Juni 2026", "APPROVED (Sistem Digital)"),
        ("Lina M.", "Information Security Manager", "24 Juni 2026", "APPROVED (Sistem Digital)")
    ]
    build_table(doc, headers_appr, rows_appr, [1.5, 2.0, 1.2, 1.8])
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # GLOSSARY
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Glosarium Istilah & Singkatan", level=1)
    headers_glos = ["Istilah / Singkatan", "Definisi / Penjelasan Singkat"]
    rows_glos = [
        ("FMSP", "Facility Management Service Platform. Platform terpadu manajemen fasilitas Lintasarta."),
        ("BRD", "Business Requirements Document. Dokumen yang menerangkan kebutuhan fungsional bisnis organisasi."),
        ("SRS", "Software Requirements Specification. Dokumen spesifikasi kebutuhan perangkat lunak formal (IEEE 830)."),
        ("RBAC", "Role-Based Access Control. Metode pembatasan hak akses sistem berdasarkan peran pengguna."),
        ("RAB", "Rencana Anggaran Biaya. Rencana alokasi anggaran tahunan divisi/departemen operasional."),
        ("WO", "Work Order. Dokumen perintah kerja penanganan komplain kerusakan fasilitas fisik."),
        ("PM", "Preventive Maintenance. Pemeliharaan preventif terjadwal untuk mendeteksi dini kerusakan aset vital."),
        ("SMK3", "Sistem Manajemen Keselamatan dan Kesehatan Kerja. Prosedur standar pengelolaan K3 gedung."),
        ("SLF", "Sertifikat Layak Fungsi. Sertifikat kelayakan bangunan gedung yang diterbitkan oleh Dinas PUPR."),
        ("IMB", "Izin Mendirikan Bangunan. Izin hukum untuk mendirikan atau merenovasi gedung."),
        ("SSO", "Single Sign-On. Metode autentikasi tunggal menggunakan kredensial Google Korporat Lintasarta."),
        ("MFA / 2FA", "Multi-Factor Authentication / Two-Factor Authentication. Lapisan keamanan tambahan setelah kata sandi."),
        ("JWT", "JSON Web Token. Token terenkripsi untuk menyimpan informasi sesi pengguna di sisi client browser."),
        ("VAPID", "Voluntary Application Server Identification. Kunci kriptografi standar untuk Push Notification."),
        ("HA", "High Availability. Konfigurasi infrastruktur untuk memastikan ketersediaan layanan terus-menerus."),
        ("CNPG", "CloudNativePG. Operator Kubernetes untuk manajemen kluster PostgreSQL High Availability.")
    ]
    build_table(doc, headers_glos, rows_glos, [1.8, 4.7])
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB I: PENDAHULUAN
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Bab I: Pendahuluan", level=1)
    
    add_heading_with_bottom_border(doc, "1.1 Latar Belakang & Analisis Kesenjangan (Gap Analysis)", level=2)
    doc.add_paragraph(
        "Sebagai perusahaan infrastruktur telekomunikasi nasional, PT Aplikanusa Lintasarta mengoperasikan ratusan aset properti "
        "mulai dari kantor pusat (Jakarta), puluhan kantor regional, pusat data (Data Center), stasiun bumi satelit (Jatiluhur), "
        "hingga shelter PoP (Points of Presence) yang tersebar di wilayah terpencil. Kompleksitas manajemen seluruh fasilitas "
        "ini menuntut adanya sistem informasi terpadu.\n\n"
        "Sebelum diimplementasikannya FMSP Lintasarta, analisis kesenjangan (Gap Analysis) mengidentifikasi beberapa inefisiensi:"
    )
    
    headers_gap = ["Proses Bisnis Eksisting (As-Is)", "Masalah / Kesenjangan", "Solusi FMSP (To-Be)"]
    rows_gap = [
        ("Pengelolaan Izin Gedung (SLF, IMB) menggunakan Excel manual.", "Tanggal kedaluwarsa terlewat, risiko denda hukum PUPR dan penyegelan gedung.", "Modul Legalitas Aset dengan Reminder Alerts H-30 otomatis melalui Email & Push Notification."),
        ("Komplain kerusakan fasilitas dikirim via WhatsApp / Email personal.", "Tiket kerusakan tidak terdata, SLA perbaikan tidak terukur, teknisi tumpang tindih.", "Modul Work Order dengan alur status terstruktur, assign teknisi, SLA deadline, dan verifikasi approval."),
        ("Pemeliharaan mesin MEP (AC, UPS, Genset) dijadwalkan secara reaktif.", "Kerusakan mesin mendadak, biaya overhaul membengkak, downtime operasional tinggi.", "Modul Preventive Maintenance (PM) terjadwal otomatis dengan kalkulator jatuh tempo dinamis."),
        ("Pencatatan persediaan suku cadang gudang dilakukan pasca-penggunaan.", "Kehabisan stok material kritis di tengah perbaikan darurat.", "Modul Inventory dengan penanda Stok Minimum (Safety Stock) berwarna merah dan alert kritis."),
        ("Pencatatan kas GA tidak terhubung ke pos anggaran tahunan.", "Realisasi pengeluaran melebihi plafon (Over-Budget) tanpa kontrol pencegahan.", "Modul Keuangan terintegrasi langsung dengan tabel RAB dan fitur Over-Budget Lock otomatis.")
    ]
    build_table(doc, headers_gap, rows_gap, [2.0, 2.2, 2.3])
    
    add_heading_with_bottom_border(doc, "1.2 Tujuan Platform", level=2)
    doc.add_paragraph(
        "FMSP Lintasarta ditargetkan untuk memenuhi tujuan bisnis utama (SMART Goals) sebagai berikut:\n"
        "1. Kepatuhan Hukum 100%: Menjamin nol keterlambatan perpanjangan izin gedung operasional (SLF, IMB, izin genset).\n"
        "2. Optimalisasi SLA Perbaikan: Memastikan 95% tiket kerusakan kategori 'Critical' dan 'High' diselesaikan dalam kurun waktu kurang dari 24 jam.\n"
        "3. Efisiensi Biaya (Cost Control): Menekan tingkat kelebihan anggaran (over-budget) operasional FMS menjadi 0% melalui validasi pos RAB.\n"
        "4. Ketersediaan Fasilitas (Zero Downtime): Menurunkan angka kerusakan mesin utilitas utama (MEP) sebesar 40% melalui program perawatan preventif rutin."
    )
    
    add_heading_with_bottom_border(doc, "1.3 Ruang Lingkup Sistem & Batasan Batas (Scope Boundaries)", level=2)
    doc.add_paragraph(
        "Batasan fungsionalitas sistem dibagi secara tegas untuk menjaga fokus pengembangan:\n"
        "• Ruang Lingkup Masuk (In-Scope):\n"
        "  - Pengelolaan Aset Fisik Lintasarta nasional, regional, dan lokal.\n"
        "  - Penjadwalan pemeliharaan preventif, manajemen gudang, program SMK3, manajemen vendor & kontrak kerja sama.\n"
        "  - Manajemen SDM internal (teknisi, shift satpam), manajemen anggaran RAB Lintasarta GA.\n"
        "  - Alur persetujuan perbaikan Work Order, fitur chat asisten AI lokal, sistem audit log infosec.\n"
        "• Ruang Lingkup Keluar (Out-of-Scope):\n"
        "  - Pembayaran invoice keuangan secara langsung (sistem hanya mencatat data transaksi, modul transfer bank/gateway di luar FMSP).\n"
        "  - Pembelian suku cadang langsung ke e-commerce (pengadaan barang mengikuti sistem e-procurement korporat Lintasarta terpisah).\n"
        "  - Pendaftaran kehadiran satpam menggunakan GPS geofencing (sistem FMSP hanya mengatur plotting pos shift satpam)."
    )
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB II: BUSINESS REQUIREMENTS DOCUMENT (BRD)
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Bab II: Business Requirements Document (BRD)", level=1)
    
    add_heading_with_bottom_border(doc, "2.1 Aturan Bisnis & Logika Khusus (Business Rules)", level=2)
    doc.add_paragraph(
        "FMSP menerapkan aturan bisnis yang ketat untuk mengendalikan operasional secara otomatis. Detail aturan tersebut adalah:"
    )
    
    add_heading_with_bottom_border(doc, "2.1.1 Aturan Pengingat Dokumen & Vendor Contract", level=3)
    doc.add_paragraph(
        "Setiap dokumen legalitas properti (SLF/IMB) dan kontrak kerja sama vendor yang terdaftar di sistem memiliki atribut `expiryDate` (Tanggal Kedaluwarsa) dan `picEmail` (Surel Penanggung Jawab).\n"
        "• Sistem mengeksekusi scheduler harian setiap pukul 02:00 WIB untuk memindai rekam dokumen.\n"
        "• Sisa Hari Kepatuhan (Days to Expiry) dihitung dengan formula: \n"
        "  DaysRemaining = expiryDate - CurrentDate\n"
        "• Status kepatuhan (`complianceStatus`) diubah secara otomatis:\n"
        "  - Jika DaysRemaining > 30 hari, status = 'Valid' (Hijau)\n"
        "  - Jika 0 <= DaysRemaining <= 30 hari, status = 'Warning' (Kuning)\n"
        "  - Jika DaysRemaining < 0 hari, status = 'Expired' (Merah)\n"
        "• Trigger Notifikasi Email dikirimkan ke `picEmail` pada H-30, H-14, H-7, H-3, H-1, dan hari H expired."
    )
    
    add_heading_with_bottom_border(doc, "2.1.2 Aturan Stok Aman Gudang (Safety Stock Calculation)", level=3)
    doc.add_paragraph(
        "Untuk mencegah kekosongan material suku cadang kritis, setiap item gudang (`InventoryItem`) wajib memiliki nilai batas:\n"
        "  `minQty` (Stok Minimum / Safety Stock)\n"
        "  `maxQty` (Stok Maksimum)\n"
        "• Sistem secara real-time membandingkan kuantitas stok riil (`qty`): \n"
        "  - Jika qty <= minQty, tampilkan baris tabel berwarna Merah Menyala dan kirim notifikasi berkategori Kritis ke tim logistik.\n"
        "  - Jika minQty < qty <= (minQty + 10% dari minQty), tampilkan baris tabel berwarna Kuning (Peringatan Reorder).\n"
        "• Jumlah pemesanan ulang optimal (Reorder Quantity) direkomendasikan dengan formula:\n"
        "  ReorderQty = maxQty - qty"
    )
    
    add_heading_with_bottom_border(doc, "2.1.3 Aturan Penguncian Anggaran Keuangan (RAB Over-budget Lock)", level=3)
    doc.add_paragraph(
        "Sistem mencegah pemborosan kas operasional dengan membatasi transaksi keuangan di atas plafon anggaran tahunan.\n"
        "• Setiap departemen GA memiliki entri di tabel `RabBudget` (`allocatedAmount` = total plafon tahunan, `spentAmount` = realisasi pengeluaran berjalan).\n"
        "• Ketika transaksi pengeluaran (`AccountingTransaction`) senilai `amount` diinput:\n"
        "  SisaPagu = allocatedAmount - spentAmount\n"
        "  - Jika amount <= SisaPagu, transaksi disetujui. Nilai `spentAmount` diubah secara atomik: `spentAmount = spentAmount + amount`.\n"
        "  - Jika amount > SisaPagu, sistem wajib menolak transaksi (Rollback). Tampilkan pesan kesalahan 'Transaksi Ditolak: Plafon RAB Departemen Tidak Mencukupi (Over-Budget Lock)'."
    )
    
    add_heading_with_bottom_border(doc, "2.1.4 Aturan Alur Approval Status Tiket Work Order", level=3)
    doc.add_paragraph(
        "Penanganan kerusakan mengikuti mesin status (State Machine) transisi tiket kerja berikut:\n"
        "1. Open: Tiket dibuat oleh User pelapor kerusakan ( reportedBy ). Status approval kosong.\n"
        "2. In Progress: Admin Pusat / Regional menunjuk teknisi pelaksana ( assignedTo ). Tiket diproses.\n"
        "3. Resolved: Teknisi menyelesaikan perbaikan dan mengunggah foto penyelesaian. Status berubah ke Resolved. Sistem memicu permintaan review ke Manager FMS.\n"
        "4. Closed (Final): Manager FMS menyetujui hasil perbaikan. Status tiket berubah menjadi Closed. Record dikunci dari perubahan.\n"
        "5. Rejected: Manager FMS menolak hasil perbaikan (mengisi field `rejectedReason`). Status tiket kembali menjadi In Progress agar dikerjakan ulang oleh teknisi."
    )
    
    add_heading_with_bottom_border(doc, "2.2 Matriks Hak Akses Pengguna & Filter Wilayah", level=2)
    doc.add_paragraph(
        "FMSP menerapkan pembatasan data regional berdasarkan wilayah kerja staf (region scope) yang disimpan pada field `region` di tabel User. "
        "Pengecekan wilayah wajib diterapkan pada query tingkat database (Prisma middleware / query filter):"
    )
    
    headers_filter = ["Role Pengguna", "Filter Data Region", "Filter Data Lokasi / Gedung", "Kemampuan Operasi"]
    rows_filter = [
        ("superadmin", "Tidak ada (Bypass)", "Tidak ada (Bypass)", "Full CRUD semua data nasional."),
        ("manager_fms", "Tidak ada (Bypass)", "Tidak ada (Bypass)", "Read-only nasional + otorisasi Approval."),
        ("admin_pusat", "Tidak ada (Bypass)", "Tidak ada (Bypass)", "CRUD operasional tingkat nasional."),
        ("admin_regional", "Ya (Hanya region penugasan)", "Tidak ada", "CRUD operasional di region terkait. Data region lain diblokir."),
        ("admin_lokasi", "Ya (Hanya region penugasan)", "Ya (Hanya gedung penugasan)", "CRUD operasional terbatas pada satu gedung terkait."),
        ("user", "Ya (Hanya region penugasan)", "Ya (Hanya gedung penugasan)", "Read-only fasilitas gedung penugasan + Create WO.")
    ]
    build_table(doc, headers_filter, rows_filter, [1.5, 1.8, 1.8, 1.4])
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB III: SOFTWARE REQUIREMENTS SPECIFICATION (SRS)
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Bab III: Software Requirements Specification (SRS)", level=1)
    
    add_heading_with_bottom_border(doc, "3.1 Persyaratan Antarmuka Eksternal (External Interfaces)", level=2)
    doc.add_paragraph(
        "• Antarmuka Pengguna (User Interface): Desain web premium menggunakan Tailwind CSS dengan layout responsive (Desktop, Tablet, Mobile). "
        "Mendukung tema alih Dark/Light Mode. Menggunakan Lucide React untuk ikonografi dan Recharts untuk visualisasi diagram interaktif.\n"
        "• Antarmuka Perangkat Lunak (Software Interface): Hubungan ke API Google OAuth untuk SSO, dan database PostgreSQL 15 via Prisma ORM.\n"
        "• Antarmuka Komunikasi (Communication Interface): Integrasi ke server Outlook SMTP korporat Lintasarta untuk relay surel pengingat. "
        "Koneksi SSE (Server-Sent Events) untuk notifikasi real-time, serta Push API (VAPID) untuk notifikasi web browser melayang."
    )
    
    add_heading_with_bottom_border(doc, "3.2 Persyaratan Fungsional (Functional Requirements)", level=2)
    doc.add_paragraph(
        "Berikut adalah tabel rincian spesifikasi kebutuhan fungsional sistem FMSP Lintasarta:"
    )
    
    headers_req = ["ID Persyaratan", "Nama Fitur", "Prosedur Input / Trigger", "Pemrosesan & Aturan", "Output / Hasil Akhir"]
    rows_req = [
        ("REQ-001", "Login Corporate SSO", "Klik tombol 'Sign In with Google' di halaman /login.", "Google ID Token diverifikasi. Email wajib diakhiri '@lintasarta.co.id'.", "Token JWT disimpan di HttpOnly Cookie, user diarahkan ke dashboard."),
        ("REQ-002", "Brute Force Lockout", "Pengguna gagal memasukkan kata sandi 5 kali berturut-turut.", "Akun dikunci. Field `lockedUntil` diubah ke waktu saat ini + 15 menit.", "Form login dinonaktifkan untuk email tersebut dengan pesan error Lockout."),
        ("REQ-003", "Aset CRUD & QR Code", "Admin menginput nama aset, tipe, gedung, tanggal beli, nilai buku.", "Record dibuat di DB. Sistem generate QR Code unik berbasis CUID aset.", "Record aset disimpan, tombol 'Print Label' berlogo Lintasarta siap diunduh."),
        ("REQ-004", "Upload Izin Gedung 100MB", "User memilih file scan SLF/IMB berukuran hingga 100MB di form perizinan.", "API Endpoint `/api/upload` memverifikasi ukuran berkas dan menyimpan file di storage.", "Dokumen legal terunggah, URL file tersimpan di database."),
        ("REQ-005", "Auto-Reminder Legal", "Scheduler harian (cron) berjalan otomatis pukul 02:00 WIB.", "Membandingkan tanggal expiry dengan hari ini. Kirim notifikasi jika H-30.", "Email notifikasi terkirim ke PIC, alert peringatan kuning muncul di dashboard."),
        ("REQ-006", "PM Schedule Due Date", "Teknisi mengubah status PM menjadi Selesai.", "Sistem mengubah `lastPerformed` ke hari ini, dan `nextDue` ke hari ini + interval.", "Jadwal PM berikutnya dibuat otomatis di tabel `maintenance_schedules`."),
        ("REQ-007", "Safety Stock Gudang", "Kuantitas material berkurang saat digunakan untuk perbaikan.", "Sistem mencocokkan `qty` dengan `minQty`. Jika qty <= minQty, set status alert.", "Baris item gudang berwarna merah menyala, alert kritis logistik dipicu."),
        ("REQ-008", "RAB Budget Enforcement", "User menginput pengeluaran operasional baru.", "Sistem mengevaluasi sisa saldo pos RAB. Penguncian jika melebihi plafon.", "Transaksi ditolak dengan pesan Over-Budget Lock, saldo RAB aman."),
        ("REQ-009", "Work Order State Machine", "Manager FMS mengkaji tiket WO berstatus Resolved.", "Manager menekan Approve (WO ditutup) atau Reject (WO kembali ke teknisi).", "Status WO berubah ke Closed (Final) atau kembali ke In Progress."),
        ("REQ-010", "AI Copilot Sop Guide", "User mengirim pertanyaan mengenai SOP perbaikan Genset ke widget chat.", "Route API memanggil Ollama. Jika timeout >1.2 detik, gunakan pencarian offline.", "Widget menampilkan SOP perbaikan atau panduan manual kontekstual.")
    ]
    build_table(doc, headers_req, rows_req, [1.0, 1.2, 1.4, 1.6, 1.3])
    
    add_heading_with_bottom_border(doc, "3.3 Persyaratan Non-Fungsional (Non-Functional Requirements)", level=2)
    
    headers_nfr = ["Kategori NFR", "Kode", "Persyaratan Detil Standar Industri"]
    rows_nfr = [
        ("Security (Keamanan)", "NFR-SEC-001", "Seluruh koneksi lalu lintas data wajib dienkripsi penuh menggunakan SSL/TLS (HTTPS). Kata sandi disimpan menggunakan bcrypt hashing."),
        ("Security (Keamanan)", "NFR-SEC-002", "Sesi login menggunakan JWT Token di HTTP-only cookie dengan perlindungan strict CORS whitelisting untuk mencegah CSRF."),
        ("Performance (Kinerja)", "NFR-PER-001", "Waktu respon halaman dashboard utama (pemuatan metrik KPI) tidak boleh melebihi 1.5 detik pada jaringan koneksi normal."),
        ("Performance (Kinerja)", "NFR-PER-002", "Widget AI Copilot harus mengaktifkan pencarian kata kunci offline (fallback) jika server Ollama lokal tidak merespons dalam 1.2 detik."),
        ("Reliability (Keandalan)", "NFR-REL-001", "Aplikasi harus pulih secara otomatis (auto-restart) jika terjadi crash server, dengan pemulihan koneksi database singleton Prisma."),
        ("Availability (Ketersediaan)", "NFR-AVA-001", "Tingkat ketersediaan layanan sistem FMSP minimal sebesar 99.9% (maksimal downtime tidak terencana 8 jam per tahun)."),
        ("Usability (Kemudahan)", "NFR-USA-001", "Antarmuka sistem harus memenuhi standar desain UI modern, responsif mobile, mendukung Dark/Light Mode, dan ramah keyboard navigasi.")
    ]
    build_table(doc, headers_nfr, rows_nfr, [1.5, 1.0, 4.0])
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB IV: ARSITEKTUR FULL STACK & SPESIFIKASI TEKNIS
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Bab IV: Arsitektur Full Stack & Spesifikasi Teknis", level=1)
    
    add_heading_with_bottom_border(doc, "4.1 Detail Spesifikasi Kamus Data Database (PostgreSQL Schema)", level=2)
    doc.add_paragraph(
        "Kamus data ini menerangkan secara rinci nama tabel, kolom, jenis tipe data, nullable, serta deskripsi fungsional bisnis "
        "yang tertuang pada file schema Prisma:"
    )
    
    # 1. Tabel users
    add_heading_with_bottom_border(doc, "4.1.1 Tabel users", level=3)
    headers_t_user = ["Kolom (Field)", "Tipe Data", "Nullable / Key", "Keterangan / Aturan Bisnis"]
    rows_t_user = [
        ("id", "String (CUID)", "No / PK", "ID unik pengguna."),
        ("email", "String", "No / Unique", "Surel korporat staf (akhiran @lintasarta.co.id)."),
        ("name", "String", "No", "Nama lengkap staf."),
        ("role", "String", "No", "Peran hak akses (superadmin, manager_fms, admin_pusat, admin_regional, admin_lokasi, user)."),
        ("passwordHash", "String", "No", "Password terenkripsi algoritma BCrypt."),
        ("isActive", "Boolean", "No (default true)", "Penanda status keaktifan akun."),
        ("mustChangePassword", "Boolean", "No (default false)", "Flag paksa ganti sandi saat masuk pertama kali."),
        ("region", "String", "Yes", "Wilayah kerja regional staf untuk filter data."),
        ("failedLoginAttempts", "Integer", "No (default 0)", "Jumlah kumulatif kegagalan login berturut-turut."),
        ("lockedUntil", "DateTime", "Yes", "Masa waktu penahanan lockout akun.")
    ]
    build_table(doc, headers_t_user, rows_t_user, [1.5, 1.2, 1.2, 2.6])
    
    # 2. Tabel password_reset_requests
    add_heading_with_bottom_border(doc, "4.1.2 Tabel password_reset_requests", level=3)
    headers_t_reset = ["Kolom (Field)", "Tipe Data", "Nullable / Key", "Keterangan / Aturan Bisnis"]
    rows_t_reset = [
        ("id", "String (CUID)", "No / PK", "ID unik pengajuan."),
        ("userId", "String", "No / FK", "Relasi ke users.id."),
        ("status", "String", "No (default pending)", "Status approval (pending | approved | rejected)."),
        ("requestedAt", "DateTime", "No (default now)", "Waktu inisiasi pengajuan."),
        ("reviewedBy", "String", "Yes", "Surel admin penyetuju/penolak."),
        ("token", "String", "Yes / Unique", "Token acak reset kata sandi.")
    ]
    build_table(doc, headers_t_reset, rows_t_reset, [1.5, 1.2, 1.2, 2.6])
    
    # 3. Tabel assets
    add_heading_with_bottom_border(doc, "4.1.3 Tabel assets", level=3)
    headers_t_asset = ["Kolom (Field)", "Tipe Data", "Nullable / Key", "Keterangan / Aturan Bisnis"]
    rows_t_asset = [
        ("id", "String (CUID)", "No / PK", "ID unik aset fisik."),
        ("name", "String", "No", "Nama perangkat/fasilitas (Genset Perkins 500kVA)."),
        ("type", "String", "No", "Kategori aset (land | office | facility | vehicle)."),
        ("location", "String", "No", "Gedung penempatan fisik aset."),
        ("specs", "JSON", "No", "Spesifikasi dinamis (kapasitas, tegangan, merk)."),
        ("status", "String", "No (default good)", "Kelayakan fisik (good | warning | broken)."),
        ("bookValue", "Float", "No", "Nilai depresiasi buku aset saat ini (Rupiah)."),
        ("purchaseCost", "Float", "Yes", "Harga beli awal aset.")
    ]
    build_table(doc, headers_t_asset, rows_t_asset, [1.5, 1.2, 1.2, 2.6])

    # 4. Tabel legal_documents
    add_heading_with_bottom_border(doc, "4.1.4 Tabel legal_documents", level=3)
    headers_t_legal = ["Kolom (Field)", "Tipe Data", "Nullable / Key", "Keterangan / Aturan Bisnis"]
    rows_t_legal = [
        ("id", "String (CUID)", "No / PK", "ID unik dokumen legal."),
        ("assetId", "String", "No / FK", "Tautan kepemilikan aset (Asset.id)."),
        ("title", "String", "No", "Judul berkas (Sertifikat Layak Fungsi Lintasarta Jatiluhur 2026)."),
        ("documentType", "String", "No", "Jenis berkas (slf, certificate, insurance, tax_building, dll)."),
        ("documentUrl", "String", "No", "URL lokasi file hasil upload."),
        ("expiryDate", "DateTime", "No", "Tanggal berakhir berlakunya berkas perizinan."),
        ("complianceStatus", "String", "No (default valid)", "Status keabsahan berkas (valid | warning | expired).")
    ]
    build_table(doc, headers_t_legal, rows_t_legal, [1.5, 1.2, 1.2, 2.6])
    
    # 5. Tabel work_orders
    add_heading_with_bottom_border(doc, "4.1.5 Tabel work_orders", level=3)
    headers_t_wo = ["Kolom (Field)", "Tipe Data", "Nullable / Key", "Keterangan / Aturan Bisnis"]
    rows_t_wo = [
        ("id", "String (CUID)", "No / PK", "ID unik tiket kerja."),
        ("ticketNumber", "String", "No / Unique", "Nomor tiket berformat standar: WO-2026-XXXX."),
        ("title", "String", "No", "Judul gangguan kerusakan fasilitas."),
        ("description", "String (Text)", "No", "Kronologi / deskripsi detil kerusakan."),
        ("priority", "String", "No", "Tingkat keparahan (critical | high | medium | low)."),
        ("status", "String", "No (default open)", "Tahapan status (open | in_progress | resolved | closed)."),
        ("assignedTo", "String", "Yes", "NIP teknisi pelaksana perbaikan."),
        ("reportedBy", "String", "No", "Surel staf pelapor komplain.")
    ]
    build_table(doc, headers_t_wo, rows_t_wo, [1.5, 1.2, 1.2, 2.6])

    # 6. Tabel rab_budgets & accounting_transactions
    add_heading_with_bottom_border(doc, "4.1.6 Tabel rab_budgets & accounting_transactions", level=3)
    headers_t_rab = ["Tabel", "Kolom (Field)", "Tipe Data", "Aturan Relasi / PK / FK"]
    rows_t_rab = [
        ("rab_budgets", "id", "String (CUID)", "Primary Key (PK)"),
        ("rab_budgets", "allocatedAmount", "Float", "Alokasi plafon anggaran tahunan."),
        ("rab_budgets", "spentAmount", "Float", "Akumulasi realisasi pengeluaran berjalan."),
        ("accounting_transactions", "id", "String (CUID)", "Primary Key (PK)"),
        ("accounting_transactions", "amount", "Float", "Nominal pengeluaran operasional."),
        ("accounting_transactions", "rabBudgetId", "String (FK)", "Tautan ke rab_budgets.id (N:1).")
    ]
    build_table(doc, headers_t_rab, rows_t_rab, [1.5, 1.5, 1.2, 2.3])

    add_heading_with_bottom_border(doc, "4.1.7 Tabel Lainnya (Sesuai schema.prisma)", level=3)
    doc.add_paragraph(
        "• `employees`: Menyimpan data profil karyawan sipil/teknisi dan satpam (gadaLevel, ktaNumber, KTA Expiry).\n"
        "• `inventory_items`: Menyimpan stock logistik gudang (sku, qty, minQty, maxQty, unitPrice).\n"
        "• `smk3_items`: Menyimpan checklist kelayakan sarana proteksi K3 (APAR, detektor asap, hidran).\n"
        "• `vendor_contracts`: Kontrak rekanan outsourcing pemeliharaan genset/lift.\n"
        "• `notifications`: Record trigger alerts dan notifikasi email terjadwal.\n"
        "• `audit_logs`: Log rekam audit aktivitas user (user, action, resource, details, ip).\n"
        "• `app_settings`: Pengaturan konfigurasi sistem SMTP dan switch bot AI Copilot nasional."
    )
    
    add_heading_with_bottom_border(doc, "4.2 Spesifikasi Detail Endpoint REST API Proyek", level=2)
    doc.add_paragraph(
        "Berikut adalah deskripsi lengkap format payload request dan response untuk API penentu FMSP Lintasarta:"
    )
    
    # API 1: Auth Login
    add_heading_with_bottom_border(doc, "4.2.1 Endpoint POST /api/auth/login", level=3)
    doc.add_paragraph(
        "• Request Body (JSON):\n"
        "  {\n"
        "    \"email\": \"admin.regional@lintasarta.co.id\",\n"
        "    \"password\": \"adminregional123\"\n"
        "  }\n"
        "• Response Sukses (Status 200 OK):\n"
        "  {\n"
        "    \"success\": true,\n"
        "    \"user\": { \"id\": \"cuid_1\", \"name\": \"Admin Regional\", \"role\": \"admin_regional\", \"region\": \"Sumatera\" }\n"
        "  }\n"
        "• Response Gagal (Status 401 Unauthorized / Lockout):\n"
        "  {\n"
        "    \"success\": false,\n"
        "    \"error\": \"Akun dikunci sementara selama 15 menit akibat 5x gagal login berturut-turut.\"\n"
        "  }"
    )
    
    # API 2: SSO Auth
    add_heading_with_bottom_border(doc, "4.2.2 Endpoint POST /api/auth/sso", level=3)
    doc.add_paragraph(
        "• Request Body (JSON):\n"
        "  {\n"
        "    \"idToken\": \"google_oauth_id_token_jwt_string\"\n"
        "  }\n"
        "• Response Sukses (Status 200 OK):\n"
        "  {\n"
        "    \"success\": true,\n"
        "    \"user\": { \"email\": \"bambang@lintasarta.co.id\", \"name\": \"Bambang W\", \"role\": \"manager_fms\" }\n"
        "  }\n"
        "• Response Gagal (Status 403 Forbidden - Domain Mismatch):\n"
        "  {\n"
        "    \"success\": false,\n"
        "    \"error\": \"Hanya akun Google korporat Lintasarta (@lintasarta.co.id) yang diizinkan masuk ke sistem.\"\n"
        "  }"
    )

    # API 3: Work Order Approve/Reject
    add_heading_with_bottom_border(doc, "4.2.3 Endpoint POST /api/management/workorder/[id]/approve", level=3)
    doc.add_paragraph(
        "• Request Body (JSON):\n"
        "  {\n"
        "    \"action\": \"approve\" // approve | reject\n"
        "    \"reason\": \"Hasil perbaikan AC sudah oke, kelistrikan stabil.\" // Diisi jika action=reject\n"
        "  }\n"
        "• Response Sukses (Status 200 OK):\n"
        "  {\n"
        "    \"success\": true,\n"
        "    \"message\": \"Tiket Work Order berhasil diverifikasi dan ditutup (Closed).\",\n"
        "    \"ticket\": { \"id\": \"wo_1\", \"status\": \"closed\", \"approvedBy\": \"manager@lintasarta.co.id\" }\n"
        "  }"
    )
    
    add_heading_with_bottom_border(doc, "4.3 Alur Integrasi AI Copilot & Offline Keyword Fallback", level=2)
    doc.add_paragraph(
        "Fitur asisten cerdas AI Copilot dirancang agar tetap tangguh (resilient) jika jaringan intranet Lintasarta terputus. "
        "Logika integrasi digambarkan sebagai berikut:\n"
        "1. Saat user mengetikkan prompt di form melayang, web client mengirim POST request ke `/api/chat`.\n"
        "2. API route Next.js bertindak sebagai proxy ke service Ollama lokal di URL `http://localhost:11434`.\n"
        "3. Route handler mengaktifkan timeout Promise sebesar 1.2 detik (1200ms).\n"
        "4. Jika service Ollama gagal merespon tepat waktu atau status key `ai_bot_enabled` dinonaktifkan oleh SuperAdmin:\n"
        "   - Client API mengembalikan status 200 OK dengan flag `fallback: true`.\n"
        "   - Frontend mengaktifkan script pencarian indeks kata kunci lokal (Keyword Indexer).\n"
        "   - Sistem mencocokkan kata kunci masukan (contoh: 'Genset', 'APAR', 'SLF') dengan berkas panduan SOP yang tersimpan lokal di browser cache (Help Guide cache), dan menyajikan teks solusi SOP langsung di widget."
    )
    
    add_heading_with_bottom_border(doc, "4.4 Desain High Availability (HA) & Prosedur Disaster Recovery", level=2)
    doc.add_paragraph(
        "Untuk menjamin keandalan sistem berskala nasional:\n"
        "• High Availability (HA): Database PostgreSQL di produksi direkomendasikan berjalan di atas Kubernetes menggunakan operator CloudNativePG (CNPG) yang mengelola 3 replika node database (1 primary, 2 read-replicas). PgBouncer dikonfigurasi di depan database untuk efisiensi connection pooling.\n"
        "• Backup Otomatis: Script `/app/scripts/backup-db.sh` mengeksekusi `pg_dump` harian pada jam 01:00 WIB dini hari. Hasil dump terkompresi dikirim ke S3-compatible cold storage eksternal.\n"
        "• Disaster Recovery (DR): Jika database utama hancur, tim DevOps Lintasarta menjalankan script restore dengan perintah:\n"
        "  `pg_restore -h <db_host> -U <db_user> -d <db_name> --clean --no-owner /path/to/backup_file.dump`\n"
        "Proses pemulihan ini diuji berkala setiap semester untuk memastikan Recovery Time Objective (RTO) < 2 jam dan Recovery Point Objective (RPO) < 24 jam."
    )
    
    # Save documents
    doc.save("BRD_SRS_dan_Arsitektur_FMSP_Lintasarta.docx")
    
    # Copy to public/docs
    os.makedirs("public/docs", exist_ok=True)
    doc.save("public/docs/BRD_SRS_dan_Arsitektur_FMSP_Lintasarta.docx")
    
    print("Detailed Document generated successfully!")

if __name__ == "__main__":
    generate_detailed_document()
