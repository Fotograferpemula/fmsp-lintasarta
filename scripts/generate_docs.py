import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_bottom_border(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(23, 105, 255) # Lintasarta Blue
        run.bold = True
        
        pPr = heading._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # 1.5 pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1769FF')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(13, 79, 204) # Darker Blue
        run.bold = True
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(60, 60, 60)
        run.bold = True
        run.italic = True
        
    return heading

def create_user_manual():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # COVER PAGE
    # Logo
    logo_path = "public/docs/screenshots/logo.png"
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(40)
        logo_run = logo_p.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.8))

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_run = title_p.add_run("BUKU PUTIH & USER MANUAL")
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(23, 105, 255)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    sub_run = sub_p.add_run("FACILITY MANAGEMENT SERVICE PLATFORM (FMSP) LINTASARTA")
    sub_run.font.size = Pt(14)
    sub_run.bold = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run("―" * 40)
    line_run.font.color.rgb = RGBColor(23, 105, 255)
    
    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p.paragraph_format.space_before = Pt(60)
    version_run = version_p.add_run("Versi 2.0 (Edisi Pembaruan AI & Dokumentasi Halaman Lengkap)\nJuni 2026")
    version_run.font.size = Pt(11)
    version_run.font.color.rgb = RGBColor(120, 120, 120)
    
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.paragraph_format.space_before = Pt(100)
    org_run = org_p.add_run("PT Aplikanusa Lintasarta\nDivisi Facility Management & Keamanan Informasi")
    org_run.bold = True
    org_run.font.size = Pt(12)
    org_run.font.color.rgb = RGBColor(23, 105, 255)
    
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    add_heading_with_bottom_border(doc, "Daftar Isi", level=1)
    
    def add_toc_item(text, indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10)
        if indent == 0:
            r.bold = True
            
    add_toc_item("1. Pendahuluan & Gambaran Umum", 0)
    add_toc_item("1.1 Latar Belakang Masalah", 1)
    add_toc_item("1.2 Tujuan Platform", 1)
    add_toc_item("2. Keamanan Sistem & Arsitektur", 0)
    add_toc_item("2.1 Autentikasi Single Sign-On (SSO) & MFA", 1)
    add_toc_item("2.2 Role-Based Access Control (RBAC) & Filter Region", 1)
    add_toc_item("2.3 Manajemen Brute Force & Audit Logging", 1)
    add_toc_item("2.4 Manajemen Pengguna & Reset Password", 1)
    add_toc_item("2.5 Pengaturan Data Master (Master Data Configuration)", 1)
    add_toc_item("2.6 Log Audit Sistem (System Audit Log)", 1)
    add_toc_item("3. Panduan Operasional Modul Fase 1", 0)
    add_toc_item("3.1 Dashboard (Overview)", 1)
    add_toc_item("3.2 Manajemen Aset Fisik & Cetak QR Code", 1)
    add_toc_item("3.3 Dokumen Legalitas Aset", 1)
    add_toc_item("3.4 Kotak Pengingat & Alerts (Reminder & Alerts)", 1)
    add_toc_item("3.5 Pemeliharaan Preventif (Preventive Maintenance - PM)", 1)
    add_toc_item("3.6 Inventaris Gudang & Safety Stock", 1)
    add_toc_item("3.7 Sistem Manajemen K3 (SMK3 Safety & Insiden)", 1)
    add_toc_item("3.8 Vendor & Manajemen Kontrak", 1)
    add_toc_item("4. Panduan Operasional Modul Fase 2 (Advance Menu)", 0)
    add_toc_item("4.1 HRD & Security Shift", 1)
    add_toc_item("4.2 Manajemen Keuangan & RAB Terintegrasi", 1)
    add_toc_item("4.3 Modul Work Order (WO) & Approval Alur Kerusakan", 1)
    add_toc_item("4.4 Analytics & Visualisasi Data", 1)
    add_toc_item("5. Sistem Asisten AI & Bantuan Kontekstual", 0)
    add_toc_item("5.1 Bantuan Kontekstual Drawer (Help Guide)", 1)
    add_toc_item("5.2 Asisten AI Copilot Lokal (Ollama Integration)", 1)
    add_toc_item("5.3 Menu Konfigurasi AI & Uji Konektivitas", 1)
    add_toc_item("6. Peningkatan Kapasitas & Standar Upload", 0)
    add_toc_item("6.1 Batas Ukuran 100MB per Dokumen", 1)
    add_toc_item("6.2 Antarmuka Tombol Unggah Langsung", 1)
    add_toc_item("6.3 Menu Unduh Dokumen Pendukung & Buku Manual", 1)
    
    doc.add_page_break()
    
    # 1. PENDAHULUAN
    add_heading_with_bottom_border(doc, "1. Pendahuluan & Gambaran Umum", level=1)
    doc.add_paragraph(
        "Facility Management Service Platform (FMSP) Lintasarta adalah platform berbasis web terintegrasi "
        "yang dirancang untuk mengelola seluruh aspek pemeliharaan fasilitas fisik, aset kantor, "
        "dokumen legalitas, sistem inventaris gudang, program SMK3, hingga pengelolaan anggaran tahunan "
        "(RAB) dan ketenagakerjaan teknisi serta personil keamanan secara terpusat."
    )
    
    add_heading_with_bottom_border(doc, "1.1 Latar Belakang Masalah", level=2)
    doc.add_paragraph(
        "Sebelum diimplementasikannya FMSP, pengelolaan fasilitas di PT Aplikanusa Lintasarta dilakukan secara "
        "manual atau terfragmentasi melalui berbagai aplikasi spreadsheet yang terpisah. Hal ini menimbulkan "
        "beberapa masalah kritis:\n"
        "1. Keterlambatan Perpanjangan Dokumen: Dokumen perizinan gedung (seperti IMB, SLF, asuransi, PBB) seringkali "
        "terlambat diperpanjang karena tidak adanya sistem pengingat otomatis.\n"
        "2. Kerusakan Aset Tidak Terlacak: Tidak adanya pelaporan tiket kerusakan yang sistematis (Work Order) "
        "mengakibatkan penundaan perbaikan fasilitas vital (seperti AC Presisi di Ruang Server, Genset Backup, Panel Listrik).\n"
        "3. Over-budget Anggaran: Pihak General Affairs (GA) dan Maintenance mengalami kesulitan melacak penggunaan "
        "anggaran operasional secara riil dibandingkan pagu Rencana Anggaran Biaya (RAB) tahunan."
    )
    
    add_heading_with_bottom_border(doc, "1.2 Tujuan Platform", level=2)
    doc.add_paragraph(
        "FMSP Lintasarta hadir untuk memecahkan masalah tersebut dengan menyediakan:\n"
        "• Sistem Pengingat Dokumen (Reminder Alert): Mengirimkan alert notifikasi email otomatis H-30 sebelum masa berlaku dokumen habis.\n"
        "• Modul Preventive Maintenance (PM) & Work Order: Mengotomatisasi jadwal perawatan AC, Genset, Panel Listrik, "
        "dan menyediakan alur approval tiket kerusakan yang terstruktur.\n"
        "• Kontrol Inventaris yang Akurat: Alert otomatis berwarna merah jika stok sparepart vital di gudang berada di bawah Stok Minimum.\n"
        "• Real-time Budget Tracking: Pengeluaran keuangan terhubung langsung ke pos RAB tahunan dan memotong alokasi plafon secara otomatis."
    )
    
    doc.add_page_break()
    
    # 2. KEAMANAN SISTEM & ARSITEKTUR
    add_heading_with_bottom_border(doc, "2. Keamanan Sistem & Arsitektur", level=1)
    doc.add_paragraph(
        "Aplikasi FMSP Lintasarta dienkripsi menggunakan protokol HTTPS secara penuh. Akses masuk dikendalikan "
        "melalui SSO Korporat dan sistem logging audit yang ketat."
    )
    
    add_heading_with_bottom_border(doc, "2.1 Autentikasi Single Sign-On (SSO) & MFA", level=2)
    doc.add_paragraph(
        "Akses masuk ke aplikasi FMSP Lintasarta dikendalikan secara ketat melalui integrasi Corporate SSO Gateway Lintasarta. "
        "Setelah pengguna memasukkan email korporat (@lintasarta.co.id) dan kata sandi, sistem secara otomatis meminta "
        "verifikasi 2-Factor Authentication (2FA) melalui aplikasi Microsoft Authenticator pada perangkat telepon seluler "
        "pengguna sebelum memberikan token otorisasi JWT."
    )
    
    add_heading_with_bottom_border(doc, "2.2 Role-Based Access Control (RBAC) & Filter Region", level=2)
    doc.add_paragraph(
        "Sistem menerapkan Role-Based Access Control (RBAC) yang sangat ketat untuk membedakan hak akses "
        "berdasarkan peran masing-masing karyawan. Peran yang tersedia meliputi SuperAdmin, Manager FMS, Admin Pusat, Admin Regional, Admin Lokasi, dan User/Teknisi."
    )
    
    # Table RBAC
    table = doc.add_table(rows=6, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Role"
    hdr_cells[1].text = "Manajemen Aset"
    hdr_cells[2].text = "Work Order"
    hdr_cells[3].text = "Budget & RAB"
    hdr_cells[4].text = "Admin Master"
    for cell in hdr_cells:
        set_cell_shading(cell, "1769FF")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    row_data = [
        ("SuperAdmin", "Full CRUD (Semua)", "Full CRUD + Approve", "Full CRUD", "Ya (Akses Penuh)"),
        ("Manager FMS", "Read Only (Semua)", "Approve Selesai", "Approve & Read", "Tidak"),
        ("Admin Pusat", "Full CRUD (Semua)", "Create, Update, Assign", "Read Only", "Tidak"),
        ("Admin Regional", "CRUD Region Terkait", "Create, Update, Assign", "Read Only", "Tidak"),
        ("Teknisi / User", "Read Only", "Update Status Tugas", "Tidak Ada Akses", "Tidak")
    ]
    for idx, data in enumerate(row_data):
        row = table.rows[idx + 1]
        for c_idx, text in enumerate(data):
            cell = row.cells[c_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
            if idx % 2 == 1:
                set_cell_shading(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    
    add_heading_with_bottom_border(doc, "2.3 Manajemen Brute Force & Audit Logging", level=2)
    doc.add_paragraph(
        "Untuk mencegah serangan brute force, akun pengguna dikunci selama 15 menit jika terdeteksi 5 kali gagal login berturut-turut.\n\n"
        "Setiap operasi penambahan, pengubahan, penghapusan data, pengunggahan berkas, serta login ke sistem "
        "akan dicatat secara otomatis ke dalam database untuk mendukung audit trail infosec."
    )
    
    add_heading_with_bottom_border(doc, "2.4 Manajemen Pengguna & Reset Password", level=2)
    doc.add_paragraph(
        "Halaman **Manajemen Pengguna** (tab `users` di sidebar) digunakan oleh SuperAdmin/Admin untuk menambah staf baru, "
        "mengatur kata sandi default staf, memberikan penugasan peran (role), dan menetapkan wilayah kerja (region scope). Tim Admin "
        "dapat menyetujui atau menyetel ulang (reset) sandi staf secara manual dari menu ini apabila staf mengalami kegagalan masuk."
    )
    
    # Embed user management screenshot
    img_path = "public/docs/screenshots/approval.png"
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(8)
        p_img.add_run().add_picture(img_path, width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Gambar 2.1: Menu Manajemen Pengguna & Hak Akses Hirarki")
        r_cap.font.size = Pt(8.5)
        r_cap.italic = True
        p_cap.paragraph_format.space_after = Pt(12)
 
    add_heading_with_bottom_border(doc, "2.5 Pengaturan Data Master (Master Data Configuration)", level=2)
    doc.add_paragraph(
        "Pada halaman **Admin Master Data** (tab `admin` di sidebar), SuperAdmin dapat mengelola dropdown dan opsi dinamis di "
        "seluruh aplikasi. Hal ini mencakup penambahan gedung baru (seperti Gedung TB Simatupang, Kantor Regional), kategori vendor, "
        "jenis dokumen legalitas, jenis pemeliharaan preventif, dan divisi internal departemen. Modifikasi pada menu ini langsung "
        "terefleksi di form pengisian data tanpa memerlukan pembaruan kode program."
    )

    add_heading_with_bottom_border(doc, "2.6 Log Audit Sistem (System Audit Log)", level=2)
    doc.add_paragraph(
        "Halaman **Audit Log** (tab `auditlog` di sidebar) menampilkan daftar aktivitas operasional secara komprehensif. "
        "Setiap baris mencatat identitas pengguna, tipe aktivitas (LOGIN, CREATE, UPDATE, DELETE, UPLOAD), modul yang diakses, "
        "alamat IP asal, serta penanda waktu presisi. Admin FMS dapat melakukan pencarian berbasis nama user atau memfilter tipe "
        "operasi untuk melacak sejarah perubahan data penting."
    )

    doc.add_page_break()
    
    # 3. PANDUAN FASE 1
    add_heading_with_bottom_border(doc, "3. Panduan Operasional Modul Fase 1", level=1)
    
    add_heading_with_bottom_border(doc, "3.1 Dashboard (Overview)", level=2)
    doc.add_paragraph(
        "Halaman **Overview** (tab `overview` di sidebar) bertindak sebagai pusat informasi utama (Dashboard). Halaman ini menampilkan "
        "metrik utama secara real-time termasuk total aset fisik terdaftar, nilai buku aset, tingkat persentase kepatuhan dokumen "
        "legalitas (dokumen aktif vs kadaluarsa), sisa pagu RAB FMS, dan grafik performa kondisi kelayakan fasilitas."
    )
    
    # Embed login/dashboard screenshot
    img_path = "public/docs/screenshots/login.png"
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(8)
        p_img.add_run().add_picture(img_path, width=Inches(5.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Gambar 3.1: Tampilan Dashboard Portofolio Aset FMSP Lintasarta")
        r_cap.font.size = Pt(8.5)
        r_cap.italic = True
        p_cap.paragraph_format.space_after = Pt(12)

    add_heading_with_bottom_border(doc, "3.2 Manajemen Aset Fisik & Cetak QR Code", level=2)
    doc.add_paragraph(
        "Menu utama **Aset & Perizinan** dengan sub-tab **Aset** digunakan untuk mencatat dan memantau semua aset fisik Lintasarta. "
        "Pengguna dapat memasukkan nama aset, kategori, lokasi gedung, region, tanggal pembelian, nilai buku, dan kondisi aset. "
        "Sistem secara otomatis membuat QR Code unik untuk setiap aset baru. Pengguna dapat mengeklik tombol 'Print Label' "
        "pada tabel untuk mengunduh label QR siap cetak berlogo resmi Lintasarta untuk mempermudah scanning inventaris lapangan."
    )
    
    add_heading_with_bottom_border(doc, "3.3 Dokumen Legalitas Aset", level=2)
    doc.add_paragraph(
        "Pada menu **Aset & Perizinan** dengan sub-tab **Dokumen Legalitas**, tim FMS mengelola izin-izin gedung operasional "
        "(seperti Sertifikat Layak Fungsi - SLF, IMB, Polis Asuransi Kebakaran, Pajak PBB). Di sini pengguna dapat mengunggah pindaian "
        "dokumen (mendukung berkas hingga 100MB), mencatat tanggal berakhirnya masa berlaku, serta melihat hitungan mundur hari "
        "tersisa menuju tanggal kedaluwarsa secara visual."
    )
    
    add_heading_with_bottom_border(doc, "3.4 Kotak Pengingat & Alerts (Reminder & Alerts)", level=2)
    doc.add_paragraph(
        "Halaman **Reminder & Alerts** (tab `notifications` di sidebar) menyajikan peringatan status dokumen legalitas, status "
        "inspeksi SMK3, kontrak vendor yang akan habis, dan limit persediaan gudang. Sistem secara berkala menjalankan serverless cron "
        "dan mengirim email peringatan H-30 sebelum masa berlaku habis ke email PIC Lintasarta yang bertanggung jawab. Admin juga "
        "dapat menyetel status notifikasi menjadi 'Resolved' setelah perpanjangan selesai diproses."
    )
    
    add_heading_with_bottom_border(doc, "3.5 Pemeliharaan Preventif (Preventive Maintenance - PM)", level=2)
    doc.add_paragraph(
        "Halaman **Preventive Maintenance** (tab `maintenance` di sidebar) digunakan untuk menjadwalkan pemeriksaan berkala untuk "
        "fasilitas kritis Lintasarta seperti AC Presisi Ruang Server, Genset Backup, Panel Listrik Utama (MDP), Lift, dan APAR. "
        "Pengguna dapat mengisi interval hari pemeriksaan. Setelah status inspeksi diubah menjadi 'Selesai', sistem secara "
        "otomatis mengkalkulasi dan menjadwalkan ulang tanggal jatuh tempo pemeliharaan preventif berikutnya."
    )
    
    add_heading_with_bottom_border(doc, "3.6 Inventaris Gudang & Safety Stock", level=2)
    doc.add_paragraph(
        "Halaman **Inventory** (tab `inventory` di sidebar) melacak stok suku cadang (spareparts), APD, lampu, filter AC, dan perlengkapan "
        "pemeliharaan gedung. Setiap barang memiliki nilai batas minimum (Safety Stock). Jika jumlah stok riil berkurang di bawah batas "
        "stok minimum, sistem otomatis memberikan penanda warna merah menyala pada baris tabel serta memicu alert notifikasi agar "
        "tim logistik segera melakukan pemesanan ulang (restocking)."
    )
    
    add_heading_with_bottom_border(doc, "3.7 Sistem Manajemen K3 (SMK3 Safety & Insiden)", level=2)
    doc.add_paragraph(
        "Halaman **SMK3 Safety** (tab `smk3` di sidebar) mendukung program Keselamatan dan Kesehatan Kerja. Halaman ini terbagi menjadi "
        "dua bagian:\n"
        "1. Checklist Kelayakan Alat K3: Pemantauan kelayakan kotak P3K, tabung APAR, sistem deteksi asap (smoke detector), dan hidran.\n"
        "2. Laporan Kecelakaan Kerja & Insiden: Form pencatatan kecelakaan kerja, deskripsi kronologi, klasifikasi keparahan (minor/major), "
        "serta tindakan korektif yang wajib diisi dan diunggah bukti fotonya oleh penanggung jawab keselamatan di lokasi."
    )
    
    add_heading_with_bottom_border(doc, "3.8 Vendor & Manajemen Kontrak", level=2)
    doc.add_paragraph(
        "Halaman **Vendor & Contract** (tab `vendor` di sidebar) mencatat identitas rekanan pihak ketiga (kontraktor AC, penyedia lift, "
        "outsourcing security) serta dokumen kontrak kerja sama. Kolom pencatatan mencakup nilai kontrak, nama PIC, tanggal mulai, "
        "dan tanggal berakhir. Fitur reminder otomatis akan mengirim email alert H-30 sebelum masa kerja sama berakhir untuk "
        "mempersiapkan proses review kinerja vendor atau adendum kontrak baru."
    )
    
    doc.add_page_break()
    
    # 4. PANDUAN FASE 2
    add_heading_with_bottom_border(doc, "4. Panduan Operasional Modul Fase 2 (Advance Menu)", level=1)
    
    add_heading_with_bottom_border(doc, "4.1 HRD & Security Shift", level=2)
    doc.add_paragraph(
        "Menu **HRD & Security** (tab `hrd` di sidebar) memfasilitasi General Affairs dalam mengelola staf pendukung. Menu ini terbagi atas:\n"
        "1. Sub-tab FMS Staff / Karyawan: Mengelola data karyawan teknisi sipil, kelistrikan, tata udara, lengkap dengan info sertifikasi kompetensi.\n"
        "2. Sub-tab Security Shift: Mengatur jadwal piket satpam (pagi, siang, malam), menetapkan pembagian pos jaga (Lobby Utama, Pintu Gerbang Timur, Patroli Area Parkir), serta memantau absensi harian satpam secara terintegrasi."
    )
    
    add_heading_with_bottom_border(doc, "4.2 Manajemen Keuangan & RAB Terintegrasi", level=2)
    doc.add_paragraph(
        "Menu **Keuangan** (tab `accounting` di sidebar) mengendalikan pemakaian anggaran FMS agar tidak melampaui pagu tahunan. Terbagi menjadi:\n"
        "1. Sub-tab Keuangan / Transaksi: Pencatatan debet/kredit pengeluaran operasional (seperti pembelian freon AC, sparepart APAR, perbaikan pipa).\n"
        "2. Sub-tab Rencana Anggaran Biaya (RAB): Daftar alokasi pos anggaran tahunan FMS Lintasarta. Setiap transaksi keuangan yang diinput akan memotong sisa saldo pagu RAB bersangkutan secara real-time. Jika transaksi melebihi pagu tersisa, sistem otomatis menolak input transaksi (over-budget lock)."
    )
    
    add_heading_with_bottom_border(doc, "4.3 Modul Work Order (WO) & Approval Alur Kerusakan", level=2)
    doc.add_paragraph(
        "Menu **Work Order / Ticket** (tab `workorder` di sidebar) mengelola penanganan komplain kerusakan fasilitas. Alur penanganan kerusakan diatur sistematis:\n"
        "1. Laporan Kerusakan (Open): User menginput tiket laporan kerusakan, menautkan aset terganggu, dan memilih tingkat prioritas (Low, Medium, High, Critical).\n"
        "2. Penugasan Teknisi (On Progress): Admin menugaskan teknisi pelaksana. Status berubah menjadi On Progress.\n"
        "3. Laporan Selesai (Resolved): Teknisi mengunggah foto perbaikan dan menandai status resolved.\n"
        "4. Verifikasi & Approval (Closed): Tiket resolved harus disetujui (Approve) secara resmi oleh Manager FMS atau SuperAdmin di dashboard agar status berubah menjadi Closed. Jika hasil perbaikan belum memuaskan, Manager dapat me-reject dan mengembalikan tiket ke teknisi."
    )
    
    # Embed work order screenshot
    img_path = "public/docs/screenshots/work_order.png"
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(8)
        p_img.add_run().add_picture(img_path, width=Inches(5.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Gambar 4.1: Formulir Pengisian & Pelacakan Tiket Work Order")
        r_cap.font.size = Pt(8.5)
        r_cap.italic = True
        p_cap.paragraph_format.space_after = Pt(12)
 
    add_heading_with_bottom_border(doc, "4.4 Analytics & Visualisasi Data", level=2)
    doc.add_paragraph(
        "Menu **Analytics & Charts** (tab `analytics` di sidebar) menyediakan visualisasi grafik interaktif (ditenagai Recharts) untuk pengambilan "
        "keputusan manajemen. Grafik yang disajikan mencakup perbandingan alokasi pengeluaran OPEX vs CAPEX bulanan, tingkat penyelesaian tiket "
        "Work Order per teknisi, persentase kepatuhan sertifikat kelayakan SMK3 per regional, serta grafik pergerakan stok suku cadang gudang."
    )

    # Embed analytics screenshot
    img_path = "public/docs/screenshots/analytics.png"
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(8)
        p_img.add_run().add_picture(img_path, width=Inches(5.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Gambar 4.2: Visualisasi Charts & Grafik Analisis Keuangan FMSP")
        r_cap.font.size = Pt(8.5)
        r_cap.italic = True
        p_cap.paragraph_format.space_after = Pt(12)

    doc.add_page_break()
    
    # 5. ASISTEN AI
    add_heading_with_bottom_border(doc, "5. Sistem Asisten AI & Bantuan Kontekstual", level=1)
    
    add_heading_with_bottom_border(doc, "5.1 Bantuan Kontekstual Drawer (Help Guide)", level=2)
    doc.add_paragraph(
        "Mengeklik tombol ikon bantuan 'HelpCircle' pada form dialog akan membuka Drawer Bantuan Kontekstual di sisi kanan layar "
        "yang memandu pengisian form secara rinci tanpa mengganggu input yang sedang berjalan."
    )
    
    add_heading_with_bottom_border(doc, "5.2 Asisten AI Copilot Lokal (Ollama Integration)", level=2)
    doc.add_paragraph(
        "Obrolan interaktif AI melayang di pojok kanan bawah menggunakan model Ollama lokal (Qwen 2.5) untuk menjaga keamanan data. "
        "Memiliki fitur fallback offline instan ke Keyword Search jika server AI tidak merespons dalam 1.2 detik."
    )
    
    add_heading_with_bottom_border(doc, "5.3 Menu Konfigurasi AI & Uji Konektivitas", level=2)
    doc.add_paragraph(
        "Menu khusus admin di Phase 2 (**AI Configuration**, tab `aiconfig` di sidebar) memungkinkan SuperAdmin mengaktifkan atau menonaktifkan "
        "fitur chat widget melayang secara real-time nasional. Selain itu, menu ini menyediakan tombol interaktif 'Uji Koneksi' untuk "
        "menguji latensi ping server Ollama lokal."
    )
    
    doc.add_page_break()
    
    # 6. PENINGKATAN KAPASITAS UPLOAD
    add_heading_with_bottom_border(doc, "6. Peningkatan Kapasitas & Standar Upload", level=1)
    
    add_heading_with_bottom_border(doc, "6.1 Batas Ukuran 100MB per Dokumen", level=2)
    doc.add_paragraph(
        "Batas unggahan berkas ditingkatkan menjadi 100MB per dokumen pada endpoint `/api/upload` untuk mengakomodasi file scan resolusi tinggi, "
        "dokumen IMB tebal, dan file CAD cetak biru fasilitas."
    )
    
    add_heading_with_bottom_border(doc, "6.2 Antarmuka Tombol Unggah Langsung", level=2)
    doc.add_paragraph(
        "Dilengkapi tombol 'Upload' fisik di samping field URL pada dialog Dokumen Legalitas Aset untuk mempermudah pemilihan berkas "
        "secara lokal dari perangkat pengguna."
    )
 
    add_heading_with_bottom_border(doc, "6.3 Menu Unduh Dokumen Pendukung & Buku Manual", level=2)
    doc.add_paragraph(
        "Halaman **Dokumen & Manual** (tab `docs` di sidebar) memfasilitasi administrator untuk mengunduh versi cetak (.pdf) maupun versi dapat "
        "diedit (.docx) dari Buku Putih Panduan Pengguna dan Manual Pemeliharaan Teknis FMSP Lintasarta kapan saja secara mandiri."
    )
    
    os.makedirs("public/docs", exist_ok=True)
    doc.save("public/docs/Buku_Putih_User_Manual_FMSP_Lintasarta.docx")
    print("User Manual created successfully at public/docs/Buku_Putih_User_Manual_FMSP_Lintasarta.docx")

def create_maintenance_manual():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # COVER PAGE
    # Logo
    logo_path = "public/docs/screenshots/logo.png"
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(40)
        logo_run = logo_p.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.8))

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_run = title_p.add_run("MANUAL PEMELIHARAAN SISTEM\n(MAINTENANCE MANUAL)")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(23, 105, 255)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    sub_run = sub_p.add_run("FACILITY MANAGEMENT SERVICE PLATFORM (FMSP) LINTASARTA")
    sub_run.font.size = Pt(13)
    sub_run.bold = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run("―" * 40)
    line_run.font.color.rgb = RGBColor(23, 105, 255)
    
    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p.paragraph_format.space_before = Pt(60)
    version_run = version_p.add_run("Dokumentasi Teknis, Konfigurasi, Migrasi, & Backup\nVersi 2.0 - Juni 2026")
    version_run.font.size = Pt(11)
    version_run.font.color.rgb = RGBColor(120, 120, 120)
    
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.paragraph_format.space_before = Pt(100)
    org_run = org_p.add_run("PT Aplikanusa Lintasarta\nDivisi IT Infrastructure & DevOps")
    org_run.bold = True
    org_run.font.size = Pt(12)
    org_run.font.color.rgb = RGBColor(23, 105, 255)
    
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    add_heading_with_bottom_border(doc, "Daftar Isi", level=1)
    
    def add_toc_item(text, indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10)
        if indent == 0:
            r.bold = True
            
    add_toc_item("1. Pendahuluan & Lingkup Sistem", 0)
    add_toc_item("2. Struktur File & Komponen Terkait", 0)
    add_toc_item("2.1 Struktur Proyek Next.js", 1)
    add_toc_item("2.2 Komponen Keamanan Utama", 1)
    add_toc_item("3. Konfigurasi Sistem (System Settings)", 0)
    add_toc_item("3.1 Variabel Lingkungan (.env)", 1)
    add_toc_item("3.2 Integrasi Asisten AI (Ollama)", 1)
    add_toc_item("4. Manajemen Database & Migrasi", 0)
    add_toc_item("4.1 Prisma Schema & Client Generation", 1)
    add_toc_item("4.2 Proses Migrasi Schema", 1)
    add_toc_item("5. Prosedur Backup & Pemulihan (Backup & Restore)", 0)
    add_toc_item("5.1 Prosedur Manual Backup (pg_dump)", 1)
    add_toc_item("5.2 Prosedur Manual Restore (pg_restore)", 1)
    add_toc_item("5.3 Backup Otomatis & Pemeliharaan Rutin", 1)
    add_toc_item("6. Dokumentasi Relasi Tabel Database", 0)
    add_toc_item("6.1 Skema Hubungan Entitas (ERD)", 1)
    add_toc_item("6.2 Detail Tabel & Foreign Keys", 1)
    
    doc.add_page_break()
    
    # 1. PENDAHULUAN
    add_heading_with_bottom_border(doc, "1. Pendahuluan & Lingkup Sistem", level=1)
    doc.add_paragraph(
        "Buku Panduan Pemeliharaan ini dibuat khusus untuk administrator sistem, tim IT Infrastructure, "
        "dan DevOps PT Aplikanusa Lintasarta. Buku ini berisi instruksi teknis yang diperlukan untuk mengelola, "
        "mengonfigurasi, melakukan pencadangan data (backup), melakukan migrasi skema database, serta memelihara "
        "infrastruktur aplikasi Facility Management Service Platform (FMSP) Lintasarta agar tetap berjalan optimal "
        "dengan tingkat ketersediaan tinggi."
    )
    
    # 2. STRUKTUR FILE & KOMPONEN
    add_heading_with_bottom_border(doc, "2. Struktur File & Komponen Terkait", level=1)
    doc.add_paragraph(
        "Aplikasi FMSP Lintasarta menggunakan struktur standar Next.js App Router (TypeScript). "
        "Berikut adalah berkas dan direktori penting terkait pemeliharaan sistem:"
    )
    
    # Table Files
    table_files = doc.add_table(rows=8, cols=2)
    table_files.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_f = table_files.rows[0].cells
    hdr_f[0].text = "File / Direktori"
    hdr_f[1].text = "Fungsi & Deskripsi Teknis"
    for cell in hdr_f:
        set_cell_shading(cell, "1769FF")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    files_data = [
        ("prisma/schema.prisma", "Definisi skema database PostgreSQL dan konfigurasi Prisma client ORM."),
        ("src/lib/db.ts", "Inisialisasi koneksi Prisma Client global (mencegah kehabisan pool koneksi di dev)."),
        ("src/lib/auth-middleware.ts", "Middleware verifikasi token JWT dan proteksi autentikasi global."),
        ("src/lib/rbac.ts", "Definisi matriks peran (ROLES) dan perizinan hak akses fitur (PERMISSIONS)."),
        ("src/app/api/", "Semua endpoint REST API (backend Next.js serverless functions)."),
        ("fly.toml & Dockerfile", "Konfigurasi rilis infrastruktur containerization untuk deployment ke Fly.io."),
        ("public/docs/", "Direktori penyimpanan berkas panduan Buku Putih dan Manual Pemeliharaan (.docx / .pdf).")
    ]
    for idx, data in enumerate(files_data):
        row = table_files.rows[idx + 1]
        for c_idx, text in enumerate(data):
            cell = row.cells[c_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
            if idx % 2 == 1:
                set_cell_shading(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if c_idx == 0:
                p.runs[0].bold = True
                
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    
    # 3. KONFIGURASI SISTEM
    add_heading_with_bottom_border(doc, "3. Konfigurasi Sistem (System Settings)", level=1)
    
    add_heading_with_bottom_border(doc, "3.1 Variabel Lingkungan (.env)", level=2)
    doc.add_paragraph(
        "Aplikasi membutuhkan konfigurasi variabel lingkungan yang wajib didefinisikan baik pada file `.env` lokal "
        "maupun pada tab Secrets di Fly.io console:\n"
        "• DATABASE_URL: URL koneksi PostgreSQL utama (format: postgresql://user:password@host:port/db?schema=public).\n"
        "• JWT_SECRET: Kunci enkripsi untuk menandatangani token sesi JWT (wajib diset unik dan kuat di produksi).\n"
        "• CRON_SECRET: Kunci autentikasi untuk mengamankan endpoint cron reminder dan maintenance dari eksekusi luar.\n"
        "• APP_BASE_URL: URL root aplikasi (misalnya https://fmsp-lintasarta.fly.dev) untuk tautan reminder email.\n"
        "• SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASS: Detail server mail relay Outlook Lintasarta untuk pengiriman notifikasi."
    )
    
    add_heading_with_bottom_border(doc, "3.2 Integrasi Asisten AI (Ollama)", level=2)
    doc.add_paragraph(
        "Konfigurasi AI Copilot diatur melalui:\n"
        "• OLLAMA_BASE_URL: Default http://localhost:11434 (alamat server Ollama lokal).\n"
        "• OLLAMA_MODEL: Default qwen2.5:1.5b atau qwen2.5:3b.\n"
    )
    doc.add_paragraph(
        "Pengaturan bot asisten ini dapat dimatikan/dinyalakan secara dinamis melalui dashboard admin, "
        "dan nilainya disimpan pada tabel 'AppSetting' dengan key 'ai_bot_enabled'."
    )
    
    doc.add_page_break()
    
    # 4. MANAJEMEN DATABASE & MIGRASI
    add_heading_with_bottom_border(doc, "4. Manajemen Database & Migrasi", level=1)
    
    add_heading_with_bottom_border(doc, "4.1 Prisma Schema & Client Generation", level=2)
    doc.add_paragraph(
        "Apabila ada perubahan struktur tabel pada `prisma/schema.prisma`, tim maintenance wajib melakukan generate "
        "ulang client SDK dengan menjalankan perintah berikut pada CLI:\n"
        "  npx prisma generate"
    )
    
    add_heading_with_bottom_border(doc, "4.2 Proses Migrasi Schema", level=2)
    doc.add_paragraph(
        "Untuk menerapkan perubahan skema ke database secara aman tanpa menghapus data yang ada:\n"
        "1. Di Lingkungan Pengembangan (Development):\n"
        "  npx prisma migrate dev --name <nama_migrasi>\n"
        "2. Di Lingkungan Produksi (Fly.io/VPS):\n"
        "  npx prisma db push --accept-data-loss\n"
        "Catatan: File fly.toml telah dikonfigurasi dengan perintah rilis otomatis `release_command = \"npx prisma@6.19.3 db push --accept-data-loss\"` "
        "sehingga setiap deployment rilis baru akan otomatis menjalankan sinkronisasi skema database."
    )
    
    # 5. PROSEDUR BACKUP & PEMULIHAN
    add_heading_with_bottom_border(doc, "5. Prosedur Backup & Pemulihan (Backup & Restore)", level=1)
    
    add_heading_with_bottom_border(doc, "5.1 Prosedur Manual Backup (pg_dump)", level=2)
    doc.add_paragraph(
        "Pencadangan database PostgreSQL dilakukan secara manual dengan utilitas `pg_dump`. "
        "Jalankan perintah berikut pada terminal yang memiliki akses jaringan ke server database:\n"
        "  pg_dump -h <host_db> -U <user_db> -d <nama_db> -F c -b -v -f /path/backup/fmsp_backup_$(date +%Y%m%d_%H%M%S).dump\n\n"
        "Untuk database yang berjalan di Fly.io, Anda dapat melakukan port-forward proxy postgres lokal terlebih dahulu:\n"
        "  fly proxy 5432 -a fmsp-lintasarta-db\n"
        "Lalu jalankan pg_dump menggunakan host `localhost` and port `5432`."
    )
    
    add_heading_with_bottom_border(doc, "5.2 Prosedur Manual Restore (pg_restore)", level=2)
    doc.add_paragraph(
        "Untuk memulihkan data dari berkas cadangan (.dump) ke database target:\n"
        "  pg_restore -v -h <host_db> -U <user_db> -d <nama_db> --clean --no-owner /path/backup/fmsp_backup_file.dump\n"
        "Pilihan `--clean` akan menghapus tabel lama terlebih dahulu sebelum menimpa dengan data backup."
    )
    
    add_heading_with_bottom_border(doc, "5.3 Backup Otomatis & Pemeliharaan Rutin", level=2)
    doc.add_paragraph(
        "Sistem direkomendasikan untuk memasang cron job terjadwal pada VM Server Database Lintasarta "
        "setiap pukul 01:00 WIB dini hari untuk melakukan backup otomatis ke media cold storage eksternal.\n"
        "Contoh entri cron:\n"
        "  0 1 * * * /app/scripts/db_backup.sh >> /var/log/db_backup.log 2>&1"
    )
    
    doc.add_page_break()
    
    # 6. DOKUMENTASI RELASI TABEL
    add_heading_with_bottom_border(doc, "6. Dokumentasi Relasi Tabel Database", level=1)
    
    add_heading_with_bottom_border(doc, "6.1 Skema Hubungan Entitas (ERD)", level=2)
    doc.add_paragraph(
        "Berikut adalah visualisasi diagram hubungan entitas (ERD) dari database PostgreSQL FMSP Lintasarta "
        "yang menggambarkan keterkaitan kunci utama (Primary Key) dan kunci tamu (Foreign Key) antar tabel:"
    )

    # Embed ERD Image
    erd_path = "public/docs/erd_diagram.png"
    if os.path.exists(erd_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(8)
        p_img.add_run().add_picture(erd_path, width=Inches(5.6))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Gambar 6.1: Diagram Hubungan Entitas (ERD) Database FMSP Lintasarta")
        r_cap.font.size = Pt(8.5)
        r_cap.italic = True
        p_cap.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "1. Relasi Aset Fisik & Dokumen Legalitas (1-to-Many):\n"
        "   Tabel `Asset` dihubungkan ke tabel `LegalDocument` melalui relasi satu-ke-banyak. Hubungan didefinisikan "
        "oleh `LegalDocument.assetId` yang merujuk pada `Asset.id` (dengan aturan ON DELETE CASCADE).\n\n"
        "2. Relasi Aset Fisik & Preventive Maintenance (1-to-Many):\n"
        "   Tabel `Asset` berelasi satu-ke-banyak dengan `MaintenanceSchedule`. Kunci relasi didefinisikan oleh `MaintenanceSchedule.assetId` "
        "yang merujuk pada `Asset.id`.\n\n"
        "3. Relasi Aset Fisik & Tiket Kerusakan (1-to-Many):\n"
        "   Setiap tiket laporan kerusakan (`WorkOrder`) wajib ditautkan dengan aset fisik yang terganggu melalui "
        "relasi `WorkOrder.assetId` -> `Asset.id`.\n\n"
        "4. Relasi Transaksi Keuangan & Plafon Anggaran (Many-to-1):\n"
        "   Setiap transaksi pengeluaran keuangan (`AccountingTransaction`) harus merujuk ke salah satu pos plafon "
        "anggaran tahunan (`RabBudget`) melalui foreign key `AccountingTransaction.rabBudgetId` -> `RabBudget.id`."
    )
    
    # Table Schema Details
    add_heading_with_bottom_border(doc, "6.2 Detail Tabel & Kolom Kunci", level=2)
    
    table_rel = doc.add_table(rows=7, cols=4)
    table_rel.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_r = table_rel.rows[0].cells
    hdr_r[0].text = "Tabel Asal"
    hdr_r[1].text = "Kolom FK (Foreign Key)"
    hdr_r[2].text = "Tabel Referensi"
    hdr_r[3].text = "Jenis Kardinalitas"
    
    for cell in hdr_r:
        set_cell_shading(cell, "1769FF")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    rel_data = [
        ("LegalDocument", "assetId", "Asset", "Many-to-One (N:1)"),
        ("MaintenanceSchedule", "assetId", "Asset", "Many-to-One (N:1)"),
        ("WorkOrder", "assetId", "Asset", "Many-to-One (N:1)"),
        ("AccountingTransaction", "rabBudgetId", "RabBudget", "Many-to-One (N:1)"),
        ("AuditLog", "userId", "User", "Many-to-One (N:1)"),
        ("Notification", "userId", "User", "Many-to-One (N:1)")
    ]
    for idx, data in enumerate(rel_data):
        row = table_rel.rows[idx + 1]
        for c_idx, text in enumerate(data):
            cell = row.cells[c_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
            if idx % 2 == 1:
                set_cell_shading(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if c_idx in [0, 2]:
                p.runs[0].bold = True
                
    doc.save("public/docs/Buku_Panduan_Pemeliharaan_FMSP_Lintasarta.docx")
    print("Maintenance Manual created successfully at public/docs/Buku_Panduan_Pemeliharaan_FMSP_Lintasarta.docx")

if __name__ == "__main__":
    create_user_manual()
    create_maintenance_manual()
